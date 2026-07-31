import tempfile
import unittest
from pathlib import Path

from docx import Document

from webui.docx_reader import DocxError, read_docx_sections


class DocxReaderTests(unittest.TestCase):
    def test_preserves_paragraph_and_table_order(self):
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "ordered.docx"
            document = Document()
            document.add_paragraph("Before table")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Left"
            table.cell(0, 1).text = "Right"
            document.add_paragraph("After table")
            document.save(source)

            sections = read_docx_sections(source)

        self.assertEqual(len(sections), 1)
        self.assertLess(
            sections[0].text.index("Before table"),
            sections[0].text.index("Left | Right"),
        )
        self.assertLess(
            sections[0].text.index("Left | Right"),
            sections[0].text.index("After table"),
        )

    def test_escapes_preview_html(self):
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "escaped.docx"
            document = Document()
            document.add_paragraph("<script>alert('unsafe')</script>")
            document.save(source)

            section = read_docx_sections(source)[0]

        self.assertNotIn("<script>", section.source_html)
        self.assertIn("&lt;script&gt;", section.source_html)

    def test_rejects_invalid_docx(self):
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "invalid.docx"
            source.write_bytes(b"not a zip package")
            with self.assertRaisesRegex(DocxError, "valid DOCX"):
                read_docx_sections(source)


if __name__ == "__main__":
    unittest.main()
